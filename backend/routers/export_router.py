from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from datetime import datetime
from backend import database, models
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import io
import json

router = APIRouter(prefix="/api/export", tags=["Export"])

@router.get("/report/pdf")
def export_report_pdf(db: Session = Depends(database.get_db)):
    """Export system report as PDF"""
    try:
        users = db.query(models.User).all()
        videos = db.query(models.Video).all()
        zones = db.query(models.Zone).all()
        analyses = db.query(models.AnalysisResult).all()
        
        avg_crowd = 0
        max_crowd = 0
        if analyses:
            avg_crowd = sum(a.total_count for a in analyses) // len(analyses)
            max_crowd = max(a.total_count for a in analyses)
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(0, 102, 204),
            spaceAfter=30,
            alignment=1
        )
        
        elements.append(Paragraph("CROWD COUNT ANALYTICS - SYSTEM REPORT", title_style))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        elements.append(Spacer(1, 20))
        
        data = [
            ['Metric', 'Value'],
            ['Total Users', str(len(users))],
            ['Total Videos', str(len(videos))],
            ['Total Analyses', str(len(analyses))],
            ['Total Zones', str(len(zones))],
            ['Average Crowd Count', str(avg_crowd)],
            ['Maximum Crowd Count', str(max_crowd)],
        ]
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)
        
        return StreamingResponse(
            iter([buffer.getvalue()]),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=report-{datetime.now().strftime('%Y%m%d')}.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/report/docx")
def export_report_docx(db: Session = Depends(database.get_db)):
    """Export system report as DOCX"""
    try:
        users = db.query(models.User).all()
        videos = db.query(models.Video).all()
        zones = db.query(models.Zone).all()
        analyses = db.query(models.AnalysisResult).all()
        
        avg_crowd = 0
        max_crowd = 0
        if analyses:
            avg_crowd = sum(a.total_count for a in analyses) // len(analyses)
            max_crowd = max(a.total_count for a in analyses)
        
        doc = Document()
        doc.add_heading('CROWD COUNT ANALYTICS - SYSTEM REPORT', 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        doc.add_heading('Summary Statistics', level=1)
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'
        
        data = [
            ['Total Users', str(len(users))],
            ['Total Videos', str(len(videos))],
            ['Total Analyses', str(len(analyses))],
            ['Total Zones', str(len(zones))],
            ['Average Crowd Count', str(avg_crowd)],
            ['Maximum Crowd Count', str(max_crowd)],
        ]
        
        for i, (metric, value) in enumerate(data, 1):
            table.rows[i].cells[0].text = metric
            table.rows[i].cells[1].text = value
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            iter([buffer.getvalue()]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=report-{datetime.now().strftime('%Y%m%d')}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/crowd-data/pdf")
def export_crowd_data_pdf(data: dict):
    """Export crowd data as PDF"""
    try:
        rows = data.get('rows', [])
        if not rows:
            raise HTTPException(status_code=400, detail="No data to export")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
        elements = []
        styles = getSampleStyleSheet()
        
        elements.append(Paragraph(f"Data Export - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Heading2']))
        elements.append(Spacer(1, 12))
        
        columns = list(rows[0].keys())
        table_data = [columns]
        
        for row in rows:
            table_data.append([str(row.get(col, '')) for col in columns])
        
        table = Table(table_data, colWidths=[len(col) * 15 for col in columns])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)
        
        return StreamingResponse(
            iter([buffer.getvalue()]),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=data-export-{datetime.now().strftime('%Y%m%d')}.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/crowd-data/docx")
def export_crowd_data_docx(data: dict):
    """Export crowd data as DOCX"""
    try:
        rows = data.get('rows', [])
        if not rows:
            raise HTTPException(status_code=400, detail="No data to export")
        
        doc = Document()
        doc.add_heading('Data Export', 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        columns = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr_cells[i].text = str(col)
        
        for row in rows:
            row_cells = table.add_row().cells
            for i, col in enumerate(columns):
                row_cells[i].text = str(row.get(col, ''))
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            iter([buffer.getvalue()]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=data-export-{datetime.now().strftime('%Y%m%d')}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
